import os
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
# Define the folder path
folder_path = "Time_series_data/cards_for_2023-09-25T15:07:47Z_video"

doc = Document()

# Assign a fixed date
date_in_folder_name = "2023-09-25"
date_in_folder_name = "2023-09-25"

# Convert the folder date to a pandas Timestamp
date = pd.to_datetime(date_in_folder_name)

# Loop through each CSV file in the folder
for file_name in os.listdir(folder_path):
    if file_name.endswith(".csv"):
        # Read the CSV file into a DataFrame
        file_path = os.path.join(folder_path, file_name)
        df = pd.read_csv(file_path, header=None, names=["Date", "Price"])

        # Convert the "Date" column to Timestamps
        df["Date"] = pd.to_datetime(df["Date"])

        # Filter data for 7 days before and after the fixed date
        seven_days_before = date - pd.DateOffset(days=7)
        seven_days_after = date + pd.DateOffset(days=7)

        filtered_df = df[(df["Date"] >= seven_days_before) & (df["Date"] <= seven_days_after)]

        # Plot the data
        plt.figure(figsize=(10, 5))
        plt.plot(filtered_df["Date"], filtered_df["Price"], color='red')
        plt.title(file_name)
        plt.xlabel("Date")
        plt.ylabel("Price")
        plt.xticks(rotation=45)
        plt.tight_layout()

        # Save the plot as an image
        image_path = f"plot_{file_name.split('.')[0]}.png"
        plt.savefig(image_path)
        plt.close()

        # Add the plot image to the Word document with a specified width (e.g., 6 inches)
        doc.add_heading(file_name, level=1)
        doc.add_picture(image_path, width=Inches(6))
        doc.add_paragraph("\n")

# Save the Word document
doc.save("time_series_plots.docx")